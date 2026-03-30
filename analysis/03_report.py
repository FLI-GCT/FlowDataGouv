"""
03_report.py - Génération du rapport Word (.docx)

"Comment j'ai audité la totalité des données publiques françaises avec des Agents IA"
Structure en 3 actes, 18-22 pages.
"""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.path.insert(0, str(Path(__file__).parent))
from utils import FIGURES_DIR, OUTPUT_DIR, format_number

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x7F, 0x8C, 0x8D)


def load_stats():
    with open(FIGURES_DIR / "stats.json", "r", encoding="utf-8") as f:
        return json.load(f)


# ── Helpers Word ──────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading.append(shading.makeelement(qn("w:shd"), {qn("w:fill"): color_hex, qn("w:val"): "clear"}))


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(doc, name, width=Inches(6)):
    path = FIGURES_DIR / f"{name}.png"
    if not path.exists():
        add_para(doc, f"[Figure manquante : {name}]", italic=True, color=GRAY)
        return
    doc.add_picture(str(path), width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_stat_box(doc, number, text, bg_hex="D6EAF8"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(str(number))
    r1.font.size = Pt(28)
    r1.bold = True
    r1.font.color.rgb = NAVY
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = NAVY
    doc.add_paragraph()


def add_case_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F2F3F4")
    r1 = cell.paragraphs[0].add_run(f"CAS CONCRET : {title}")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r2 = cell.add_paragraph().add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def add_pattern_box(doc, title, text):
    """Encadré pattern technique (fond bleu très clair)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "E8F0FE")
    r1 = cell.paragraphs[0].add_run(title)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r2 = cell.add_paragraph().add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


# ── Construction du rapport ───────────────────────────────────

def build_report():
    stats = load_stats()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.different_first_page_header_footer = True

        # ── En-tête (sauf page 1) ──
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()

        # Tabulation : gauche + droite
        pf = hp.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)

        run_left = hp.add_run("Guillaume CLEMENT")
        run_left.font.name = "Calibri"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        run_tab = hp.add_run("\t\t")
        run_tab.font.size = Pt(8)

        run_right = hp.add_run("Radiographie de l'Open Data Fran\u00e7ais - Mars 2026")
        run_right.font.name = "Calibri"
        run_right.font.size = Pt(8)
        run_right.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Trait fin sous l'en-tête
        hp_pPr = hp._element.get_or_add_pPr()
        border_xml = f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>'
        hp_pPr.append(parse_xml(border_xml))

        # Tab stops : droite à la marge droite (~15cm)
        tabs_xml = (f'<w:tabs {nsdecls("w")}>'
                    f'<w:tab w:val="right" w:pos="9360"/>'
                    f'</w:tabs>')
        hp_pPr.append(parse_xml(tabs_xml))

        # ── Pied de page (sauf page 1) ──
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()

        fp_pf = fp.paragraph_format
        fp_pf.space_before = Pt(4)
        fp_pf.space_after = Pt(0)

        # Trait fin au-dessus du pied
        fp_pPr = fp._element.get_or_add_pPr()
        border_xml_top = f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>'
        fp_pPr.append(parse_xml(border_xml_top))

        # Tab stops
        fp_pPr.append(parse_xml(tabs_xml))

        run_fl = fp.add_run("https://github.com/FLI-GCT/FlowDataGouv")
        run_fl.font.name = "Calibri"
        run_fl.font.size = Pt(8)
        run_fl.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        run_ft = fp.add_run("\t\t")
        run_ft.font.size = Pt(8)

        run_fp = fp.add_run("Page ")
        run_fp.font.name = "Calibri"
        run_fp.font.size = Pt(8)
        run_fp.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Champ PAGE (numéro de page courant)
        fld_page = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        fp._element.append(fld_page)

        run_sep = fp.add_run(" / ")
        run_sep.font.name = "Calibri"
        run_sep.font.size = Pt(8)
        run_sep.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Champ NUMPAGES (nombre total de pages)
        fld_total = parse_xml(
            f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        fp._element.append(fld_total)

    rate = int(round(stats['accessibility_rate']))
    total_pub = stats.get('total_public_resources', 334000)
    coverage = int(stats.get("geo_coverage_pct", 66))

    # ══════════════════════════════════════════════════════════
    # COUVERTURE
    # ══════════════════════════════════════════════════════════

    doc.add_paragraph()
    add_para(doc, "COMMENT J'AI AUDITÉ LA TOTALITÉ", bold=True, size=28, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_para(doc, "DES DONNÉES PUBLIQUES FRANÇAISES", bold=True, size=28, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "AVEC DES AGENTS IA", bold=True, size=28, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "Bilan du projet FlowDataGouv", size=16, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_figure(doc, "fig_2_1_carte_france", width=Inches(3.5))

    add_para(doc, "Mars 2026", size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Guillaume CLEMENT", size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Analyse réalisée avec FlowDataGouv (demo-fli.fr)", size=10, italic=True, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PAGE 2 : LES 5 CHIFFRES
    # ══════════════════════════════════════════════════════════

    add_heading(doc, "Les 5 chiffres qui résument tout", level=1)

    add_stat_box(doc, f"~{rate}%",
                 f"des {format_number(total_pub)} ressources publiques sont accessibles.\n"
                 f"Les {format_number(stats['intranet_resources'])} liens internes (RIE) sont des doublons pour l'administration.")

    add_stat_box(doc, f"{stats['stale_2y_pct']}%",
                 "des jeux de données n'ont pas été mis à jour depuis plus de 2 ans.\n"
                 f"L'âge médian d'un dataset est de {stats['median_age_days']} jours (~{stats['median_age_days']//365} ans).")

    add_stat_box(doc, f"{stats['ghost_orgs']}",
                 f"organisations ({stats['ghost_orgs_pct']}%) n'ont rien publié ni mis à jour\n"
                 "depuis plus de 2 ans. Ce sont des producteurs fantômes.")

    add_stat_box(doc, f"Gini = {stats['gini']}",
                 "La production de données ouvertes est extrêmement concentrée.\n"
                 "Une poignée d'institutions produit la quasi-totalité du catalogue.")

    add_stat_box(doc, f"{format_number(stats['dead_resources'])}",
                 f"liens définitivement morts (erreur 404 ou serveur disparu).\n"
                 f"Soit {stats['dead_rate']}% des ressources publiques du catalogue.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # ACTE 1 : LE DÉFI
    # ══════════════════════════════════════════════════════════

    add_heading(doc, "Acte 1 - Le défi", level=1)

    add_heading(doc, "Le contexte", level=2)

    add_para(doc, "DataGouv publie un MCP, un connecteur pour brancher des agents IA sur les "
             "75 000 jeux de données publiques françaises. 400 000 ressources téléchargeables. "
             "J'y vois immédiatement un terrain de jeu exceptionnel : des données massives, "
             "hétérogènes, non normalisées. Exactement le type de complexité qui fait "
             "trébucher les agents IA.")

    add_para(doc, "Les objectifs : tester la visualisation humaine de données complexes, évaluer "
             "l'usage de l'IA pour enrichir la donnée source sans l'altérer, et surtout "
             "fiabiliser les réponses des agents face à une donnée imparfaite.")

    add_para(doc, "Les outils : un serveur OVH, Claude pour le développement et l'infrastructure, "
             "Mistral pour l'enrichissement sémantique et le dialogue. Pas d'équipe. "
             "Un homme et ses agents.")

    add_heading(doc, "La chaîne de fiabilité", level=2)

    add_para(doc, "Rendre un agent IA fiable sur de la donnée publique, c'est une discipline à "
             "part entière. Il faut maîtriser trois maillons, et chaque maillon fragile se "
             "paie cash dans la réponse finale.")

    add_figure(doc, "fig_0_1_chaine_fiabilite")

    add_para(doc, "Le premier maillon, la donnée source : 75 000 datasets, chacun structuré différemment, "
             "pas de normalisation entre les producteurs, une qualité variable (score de 1 à 5, médiane à 3.5). "
             "Mistral a enrichi chaque dataset (catégorie, qualité, localisation, résumé) sans altérer la source.")

    add_para(doc, "Le deuxième maillon, l'infrastructure : 400 000 ressources hébergées sur des centaines "
             "de serveurs différents. Des liens morts, des timeouts, des réseaux internes. "
             "Un scan de 8 heures pour tout tester.")

    add_para(doc, "Le troisième maillon, l'agent IA : 19 outils MCP à disposition. "
             "Et un premier constat douloureux.")

    doc.add_page_break()

    add_heading(doc, "De 30% à 2% d'erreur", level=2)

    add_para(doc, "Le premier audit des appels MCP révèle un taux d'erreur de 30%. "
             "137 erreurs sur 447 appels. Le coupable principal : l'outil resource_data, "
             "avec 82% d'échec.")

    add_para(doc, "Le diagnostic : les agents IA inventent des noms de colonnes. "
             "Ils demandent \"annee\" au lieu de \"Année\", \"tranche_effectifs\" au lieu "
             "de \"Tranche d'effectifs\". L'API retourne une erreur 400. L'agent réessaie "
             "avec un autre nom inventé. Boucle infinie.")

    add_para(doc, "La solution n'est pas de rejeter l'erreur, mais de la corriger silencieusement.")

    add_figure(doc, "fig_0_2_avant_apres_erreurs")

    add_pattern_box(doc, "PATTERN 1 : Corriger plutôt que rejeter",
                    "L'agent demande \"annee\". Le système normalise (suppression des accents, "
                    "lowercase, fuzzy match) et trouve \"Année\". L'agent reçoit ses données "
                    "sans savoir qu'il s'est trompé. Résultat : 0 erreur au lieu de 3 tentatives.")

    add_pattern_box(doc, "PATTERN 2 : Toujours retourner du contexte",
                    "Quand un filtre échoue, on retourne quand même le schéma complet. "
                    "L'agent peut corriger son filtre en un seul appel au lieu de 3-4 "
                    "tentatives aveugles.")

    add_pattern_box(doc, "PATTERN 3 : Ne jamais crasher",
                    "Une erreur dans un outil MCP ne doit jamais remonter comme une exception. "
                    "Elle doit être retournée comme du texte informatif. Un crash en cascade "
                    "annule tous les appels parallèles de l'agent.")

    add_pattern_box(doc, "PATTERN 4 : Guider dans les descriptions d'outils",
                    "Mauvais : \"Interroge les données tabulaires\". "
                    "Bon : \"Le schéma est toujours retourné. Les colonnes sont corrigées "
                    "automatiquement. Si vous ne connaissez pas les colonnes, appelez sans filtre.\"")

    add_heading(doc, "Méthodologie", level=2)

    add_para(doc, f"Périmètre : {format_number(stats['total_datasets'])} jeux de données publiés par "
             f"{format_number(stats['total_orgs'])} organisations, totalisant "
             f"{format_number(stats['total_resources'])} fichiers et ressources en ligne.")

    add_para(doc, "1. Collecte des métadonnées via l'API officielle de data.gouv.fr.")

    add_para(doc, f"2. Test de disponibilité de chaque ressource ({format_number(stats['total_resources'])} "
             "requêtes HTTP, scan de 8 heures).")

    add_para(doc, "3. Enrichissement par Mistral : catégorisation, score de qualité (1-5), "
             f"résumé, localisation géographique ({coverage}% du catalogue localisé).")

    add_para(doc, "Source des données : data.gouv.fr (API v1), audit demo-fli.fr, mars 2026. "
             "Développement avec Claude (Anthropic), enrichissement avec Mistral. "
             "Code source : github.com/FLI-GCT/FlowDataGouv", italic=True, color=GRAY)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # ACTE 2 : CE QUE LES AGENTS ONT TROUVÉ
    # ══════════════════════════════════════════════════════════

    add_heading(doc, "Acte 2 - Ce que les agents ont trouvé", level=1)

    # ── Révélation 1 ──
    add_heading(doc, "L'état réel des données publiques", level=2)

    add_para(doc, "Quand un citoyen, un chercheur ou une entreprise clique sur un lien "
             "de téléchargement sur data.gouv.fr, que se passe-t-il réellement ?", bold=True)

    add_figure(doc, "fig_1_1_verdict_global")

    add_para(doc, f"Sur les {format_number(total_pub)} ressources publiques du catalogue "
             f"(hors liens internes de l'administration), environ {rate}% sont accessibles. "
             f"{format_number(stats['dead_resources'])} liens sont définitivement morts "
             f"(erreur 404 ou serveur disparu). Les {format_number(stats['intranet_resources'])} "
             "ressources hébergées sur le réseau interministériel de l'État (RIE) sont des accès "
             "complémentaires destinés aux fonctionnaires : les données correspondantes existent "
             "déjà en accès public. Elles ne sont donc pas comptabilisées dans ce taux.")

    add_para(doc, "Un point notable : les DDT et DREAL publient souvent leurs données avec un double lien "
             "- un accès public et un accès intranet (réseau RIE). Ce dernier, destiné aux agents de l'État, "
             "apparaît dans le catalogue mais n'est pas téléchargeable par le public. Cette pratique, légitime, "
             "peut toutefois créer de la confusion pour les utilisateurs qui voient un lien \"cassé\" "
             "là où il s'agit simplement d'un accès réservé.")

    add_stat_box(doc, f"{stats['dead_rate']}%",
                 f"de liens définitivement morts ({format_number(stats['dead_resources'])} ressources)")

    add_figure(doc, "fig_1_2_fiabilite_format")

    add_para(doc, "Tous les formats ne se valent pas. Les fichiers XLSX et PDF, souvent hébergés directement "
             "sur les serveurs de data.gouv.fr, affichent des taux de disponibilité supérieurs à 90%. "
             "En revanche, les fichiers CSV et les services cartographiques (SHP, WMS) hébergés sur des serveurs "
             "externes sont nettement plus fragiles.")

    add_figure(doc, "fig_1_3_age_liens_morts")

    add_para(doc, "Le lien entre ancienneté et liens morts n'est pas linéaire. Les jeux de données "
             "de 3 mois à 2 ans affichent paradoxalement les taux de liens morts les plus élevés, "
             "tandis que les plus anciens (10 ans+) sont parmi les mieux maintenus. Ce phénomène "
             "s'explique par un biais de survie : les datasets qui traversent les années sont "
             "majoritairement ceux des grandes institutions (INSEE, IGN) qui assurent une "
             "maintenance continue de leurs liens.")

    add_figure(doc, "fig_1_5_organisations_cassees")

    add_case_box(doc, "Dijon Métropole, 19% de disponibilité",
                 "Sur les quelque 400 ressources publiées par Dijon Métropole sur data.gouv.fr, "
                 "seules 19% sont accessibles. La métropole a vraisemblablement migré ses données "
                 "vers une autre plateforme sans mettre à jour les liens sur data.gouv.fr. "
                 "Ce cas illustre un problème récurrent : quand une collectivité change de politique "
                 "de données ouvertes, les anciennes publications restent en ligne comme des coquilles vides.")

    add_case_box(doc, "Infogreffe, le registre des entreprises à 12%",
                 "Infogreffe, qui publie les données du registre du commerce et des sociétés, "
                 "n'a que 12% de ses ressources accessibles sur data.gouv.fr. Pour un service aussi "
                 "critique pour l'écosystème économique, ce taux interroge sur la pérennité des liens "
                 "entre les plateformes de données publiques et les opérateurs qui les alimentent.")

    doc.add_page_break()

    # ── Révélation 2 ──
    add_heading(doc, "La carte de France de l'Open Data", level=2)

    add_para(doc, "Tous les territoires ne sont pas égaux face à l'ouverture des données. "
             "Cette carte synthétise, pour chaque département, un indice de maturité combinant "
             "le volume de données publiées, leur qualité, leur fraîcheur et leur disponibilité effective.", bold=True)

    add_figure(doc, "fig_2_1_carte_france", width=Inches(5.5))

    add_para(doc, f"Cet indice est calculé à partir des jeux de données pour lesquels une localisation "
             f"départementale a pu être identifiée grâce à l'enrichissement par intelligence artificielle "
             f"(Mistral), soit environ {coverage}% du catalogue. Le score combine cinq dimensions : "
             f"la densité de données (rapportée à la population du département), la qualité moyenne, "
             f"la disponibilité des liens, la fraîcheur des mises à jour et la diversité thématique. "
             f"Les données à portée nationale ne sont pas représentées.")

    best = stats['best_dept']
    feminin = ["Corse", "Creuse", "Corrèze", "Charente", "Charente-Maritime", "Dordogne",
               "Drôme", "Gironde", "Guadeloupe", "Guyane", "Haute-Corse", "Haute-Garonne",
               "Haute-Loire", "Haute-Marne", "Haute-Saône", "Haute-Savoie", "Haute-Vienne",
               "Loire", "Loire-Atlantique", "Lozère", "Manche", "Marne", "Martinique",
               "Mayenne", "Meuse", "Moselle", "Nièvre", "Saône-et-Loire", "Sarthe",
               "Savoie", "Seine-Maritime", "Seine-et-Marne", "Somme", "Vendée",
               "Vienne", "Mayotte", "La Réunion"]
    if best in feminin:
        article = "La "
    elif best[0].upper() in "AEIOUYÉÈÊÎÔÙ":
        article = "L'"
    else:
        article = "Le "
    density = int(stats.get('best_dept_density', 0))
    bq = stats.get('best_dept_quality', 0)
    bh = int(stats.get('best_dept_health', 0))
    bpop = stats.get('best_dept_pop', 0)
    add_case_box(doc, f"{article}{best} en tête de l'indice",
                 f"{article}{best} arrive en tête. "
                 f"Rapporté à ses {bpop:,.0f} habitants, le département affiche "
                 f"{density} jeux de données pour 100 000 habitants, "
                 f"avec une qualité moyenne de {bq}/5 "
                 f"et {bh}% de liens fonctionnels.".replace(",", " "))

    add_figure(doc, "fig_2_2_top_bottom_departements")

    add_para(doc, "Rapporté à la population, les écarts se précisent. Certains départements ruraux "
             "affichent une densité de données remarquable grâce à des services déconcentrés "
             "de l'État (DDT, DREAL) particulièrement actifs, tandis que des départements "
             "densément peuplés restent en retrait.")

    add_figure(doc, "fig_2_3_heatmap_regions_categories")

    add_para(doc, "La heatmap ci-dessus révèle les \"déserts de données\" : certaines régions n'ont "
             "quasiment aucun jeu de données dans des thématiques pourtant essentielles. "
             "L'environnement domine largement dans toutes les régions, en grande partie grâce aux données "
             "réglementaires des DDT (plans de prévention des risques, urbanisme).")

    doc.add_page_break()

    # ── Révélation 3 ──
    add_heading(doc, "Qui publie vraiment ?", level=2)

    add_para(doc, f"Le catalogue de data.gouv.fr est alimenté par {format_number(stats['total_orgs'])} organisations. "
             "Mais la production est-elle réellement distribuée ?", bold=True)

    add_figure(doc, "fig_3_1_lorenz")

    add_para(doc, f"Le coefficient de Gini de {stats['gini']} traduit une concentration extrême, "
             "comparable à celle observée pour la distribution des revenus dans les pays les plus inégalitaires. "
             "Concrètement, une poignée d'institutions (ministères, opérateurs nationaux, grandes métropoles) "
             "produit l'écrasante majorité des données. La longue traîne est composée de milliers de petites "
             "communes et associations qui ont publié un ou deux jeux de données, souvent jamais mis à jour.")

    add_stat_box(doc, f"{stats['ghost_orgs']}",
                 f"organisations fantômes - {stats['ghost_orgs_pct']}% du total n'ont rien mis à jour depuis 2 ans")

    add_figure(doc, "fig_3_2_top30_producteurs")
    add_figure(doc, "fig_3_3_profils_organisations")

    add_para(doc, "L'analyse statistique fait émerger cinq profils distincts d'organisations, "
             "allant des grandes institutions qui produisent des centaines de jeux de données de haute qualité "
             "aux micro-producteurs qui n'ont publié qu'un ou deux fichiers.")

    add_figure(doc, "fig_3_4_organisations_fantomes")

    doc.add_page_break()

    # ── Ce que les gens cherchent vraiment ──
    usage_stats_path = FIGURES_DIR / "usage_stats.json"
    if usage_stats_path.exists():
        with open(usage_stats_path, "r", encoding="utf-8") as f:
            ustats = json.load(f)

        add_heading(doc, "Ce que les gens cherchent vraiment", level=2)

        add_para(doc, "Le catalogue de data.gouv.fr contient 75 000 jeux de données. Mais lesquels "
                 "sont réellement utilisés ? L'analyse des statistiques de consultation et de "
                 "téléchargement sur 12 mois (mars 2025 - février 2026) révèle une réalité "
                 "très concentrée.", bold=True)

        add_figure(doc, "fig_5_1_top20_usage")

        conc = ustats.get("concentration_top10_pct", 0)
        top1 = ustats.get("top1_title", "")
        add_stat_box(doc, f"{conc:.0f}%",
                     "des visites sont captées par les 10 premiers datasets.\n"
                     "L'open data français est une vitrine de 75 000 datasets\n"
                     "et une boutique de 50.")

        add_para(doc, f"Le fichier le plus consulté de France est \"{top1}\", "
                 f"avec plus d'un million de visites en 12 mois. "
                 "Derrière, on retrouve les bases de données d'entreprises, les données fiscales "
                 "et les référentiels géographiques. Ce sont des données \"evergreen\" : "
                 "consultées chaque mois, intégrées dans des applications, "
                 "indispensables au quotidien.")

        add_figure(doc, "fig_5_2_saisonnalite_usage")

        add_para(doc, "La consultation suit le rythme de la vie administrative française. "
                 "Les données fiscales et d'entreprises sont massivement consultées en début "
                 "d'année (déclarations, bilans). Les données d'urbanisme et foncières dominent "
                 "au printemps (transactions immobilières). L'été marque un creux prononcé, "
                 "avant une reprise à la rentrée de septembre. "
                 "Le mois de décembre 2025 affiche un pic inattendu (+61% par rapport à la moyenne), "
                 "porté par une explosion des consultations en \"Économie & Emploi\" "
                 "(640k visites contre 55k en novembre). Ce pic est vraisemblablement lié "
                 "aux clôtures budgétaires et aux appels de marchés publics de fin d'année.")

        add_figure(doc, "fig_5_3_evergreen")

        n_eg = ustats.get("n_evergreen", 0)
        n_eph = ustats.get("n_ephemere", 0)
        add_para(doc, f"{n_eg} datasets sont présents dans le top 50 au moins 8 mois sur 12 : "
                 f"ce sont les \"evergreen\", le socle utile du catalogue. "
                 f"{n_eph} datasets n'apparaissent qu'un ou deux mois : des pics liés à "
                 "l'actualité ou à des besoins ponctuels.")

        add_case_box(doc, "Le calendrier scolaire, 105 téléchargements pour chaque visite",
                     "Le calendrier scolaire est téléchargé 63 millions de fois par an pour "
                     "600 000 visites, soit un ratio de 105x. Pour chaque humain qui consulte "
                     "la page, 105 machines téléchargent le fichier. Intégrations calendrier, "
                     "applications mobiles, assistants IA : ce dataset illustre parfaitement "
                     "l'open data consommé par les robots au service des citoyens.")

        add_para(doc, "Le ratio téléchargements/visites distingue deux usages : les données à "
                 "ratio élevé sont consommées par des machines (APIs, intégrations automatiques), "
                 "celles à ratio faible sont consultées par des humains qui explorent. "
                 "L'open data français sert autant les robots que les citoyens.")

        add_para(doc, "Les ratios de téléchargement suggèrent que les développeurs et intégrateurs "
                 "représentent la majorité des usages (calendrier scolaire, jours fériés, "
                 "SIRENE, codes postaux), suivis des professionnels de l'immobilier (DVF, "
                 "cadastre) et des citoyens qui consultent ponctuellement des données de "
                 "référence.")

    doc.add_page_break()

    # ── Révélation 4 ──
    add_heading(doc, "La promesse et la réalité", level=2)

    add_para(doc, "Les producteurs de données s'engagent sur une fréquence de mise à jour : "
             "quotidienne, hebdomadaire, mensuelle, annuelle. "
             "Cet engagement est-il respecté ?", bold=True)

    add_figure(doc, "fig_4_1_promesse_realite")

    add_para(doc, "Le constat est sans appel. Des jeux de données déclarés en mise à jour \"quotidienne\" "
             "n'ont pas été modifiés depuis plus de deux ans en médiane. Même les données déclarées \"annuelles\" "
             "ne respectent pas leur engagement dans la majorité des cas. "
             "67% des jeux de données n'ont aucune fréquence déclarée.")

    add_stat_box(doc, f"{stats['stale_2y_pct']}%",
                 "des jeux de données n'ont pas été modifiés depuis plus de 2 ans")

    add_figure(doc, "fig_4_2_courbe_survie")

    add_para(doc, "Cette courbe montre le pourcentage du catalogue qui n'a pas été mis à jour "
             "depuis un seuil donné. Environ 70% des jeux de données n'ont pas été touchés depuis "
             "plus d'un an, et un tiers du catalogue est figé depuis plus de 5 ans. "
             "Le stock de données croît, mais la maintenance ne suit pas.")

    add_figure(doc, "fig_4_3_qualite_vs_popularite")

    add_para(doc, "La corrélation entre qualité et popularité est faible. "
             "Certains jeux de données de faible qualité sont massivement téléchargés "
             "(effet de réseau, intégrations automatiques), "
             "tandis que des données de haute qualité restent confidentielles faute de visibilité.")

    add_figure(doc, "fig_4_4_croissance_historique")

    add_para(doc, "Le catalogue croît régulièrement depuis 2013, avec une accélération notable à partir de 2020. "
             "La Loi pour une République numérique (2016) a marqué un tournant, "
             "mais la croissance la plus forte est récente.")

    add_figure(doc, "fig_4_5_saisonnalite")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # ACTE 3 : CE QU'IL FAUT RETENIR
    # ══════════════════════════════════════════════════════════

    add_heading(doc, "Acte 3 - Ce qu'il faut retenir", level=1)

    add_heading(doc, "Recommandations pour la donnée publique", level=2)

    add_heading(doc, "Pour la plateforme data.gouv.fr", level=3)

    add_para(doc, "1. Mettre en place un indicateur de santé automatique pour chaque jeu de données, "
             "testé mensuellement, et affiché publiquement. Un dataset dont les liens sont morts depuis "
             "6 mois devrait être signalé comme tel.")

    add_para(doc, "2. Identifier et contacter les organisations fantômes. "
             f"{stats['ghost_orgs']} producteurs n'ont rien publié depuis 2 ans : "
             "soit leurs données ne sont plus pertinentes, soit ils ont besoin d'accompagnement.")

    add_para(doc, f"3. Marquer visuellement les ressources internes (RIE) dans le catalogue. Les "
             f"{format_number(stats['intranet_resources'])} liens intranet sont des accès complémentaires "
             "légitimes pour l'administration, mais sans signalétique claire, ils apparaissent comme des "
             "liens cassés pour le citoyen. Un simple badge \"Accès administration\" lèverait l'ambiguïté.")

    add_heading(doc, "Pour les producteurs de données", level=3)

    add_para(doc, "1. Ne pas déclarer une fréquence de mise à jour que vous ne pouvez pas tenir. "
             "Mieux vaut déclarer \"ponctuel\" que \"quotidien\" et ne jamais mettre à jour.")

    add_para(doc, "2. Héberger les fichiers sur data.gouv.fr plutôt que sur des serveurs externes. "
             "Les ressources hébergées sur la plateforme ont un taux de disponibilité supérieur à 95%.")

    add_para(doc, "3. Supprimer les jeux de données obsolètes plutôt que de les laisser en ligne. "
             "Un catalogue propre vaut mieux qu'un catalogue volumineux.")

    add_heading(doc, "Pour les utilisateurs et développeurs", level=3)

    add_para(doc, "1. Vérifier systématiquement la date de dernière modification avant d'utiliser un dataset. "
             f"Plus de la moitié du catalogue a plus de {stats['median_age_days']//365} ans.")

    add_para(doc, "2. Privilégier les données des grandes institutions (INSEE, IGN, Météo-France) "
             "dont la qualité et la fraîcheur sont nettement supérieures à la moyenne.")

    add_para(doc, "3. Pour les intégrations automatiques, prévoir des mécanismes de fallback : "
             "les liens peuvent devenir morts à tout moment, notamment sur les serveurs externes.")

    doc.add_page_break()

    # ── Les 6 règles ──
    add_heading(doc, "Les 6 règles pour fiabiliser un agent IA sur de la donnée publique", level=2)

    add_para(doc, "Au-delà des constats sur l'open data, ce projet a permis de formaliser des "
             "règles concrètes pour quiconque veut brancher un agent IA sur de la donnée "
             "complexe. Elles s'appliquent à data.gouv.fr comme à n'importe quelle source "
             "de données hétérogène.")

    rules = [
        ("1. Enrichir avant de servir.",
         "Un catalogue brut est inexploitable par un agent. "
         "L'enrichissement par IA (catégorisation, score qualité, localisation) "
         "transforme 73 000 fiches techniques en 73 000 fiches compréhensibles."),
        ("2. Corriger silencieusement plutôt que rejeter.",
         "Un agent qui reçoit une erreur va réessayer au hasard. Un agent dont la requête "
         "est corrigée en amont reçoit ses données du premier coup."),
        ("3. Toujours fournir le contexte avec l'erreur.",
         "Quand un filtre échoue, retourner le schéma complet. L'agent peut se corriger "
         "en un appel au lieu de quatre."),
        ("4. Ne jamais crasher.",
         "Une exception dans un outil annule tous les appels parallèles. "
         "Retourner l'erreur comme du texte, pas comme une exception."),
        ("5. Le proxy avec fallback est non négociable.",
         "L'enrichissement IA tombe parfois. L'API source doit toujours être accessible en direct. "
         "Le service ne doit jamais être interrompu."),
        ("6. Monitorer les vrais usages.",
         "Les logs structurés des appels d'agents révèlent des patterns invisibles : "
         "les outils qui échouent le plus, les requêtes qui tournent en boucle, "
         "les données jamais trouvées."),
    ]
    for title, text in rules:
        add_pattern_box(doc, title, text)

    doc.add_page_break()

    # ── Et après ? ──
    add_heading(doc, "Et après ?", level=2)

    add_para(doc, "Ce projet était un terrain d'expérimentation. Les prochaines étapes pour "
             "quiconque veut aller plus loin :")

    add_para(doc, "Profiling enrichi des ressources : aller au-delà du schéma brut. Min/max "
             "pour les numériques, pourcentage de valeurs nulles, exemples de valeurs. "
             "Un agent qui sait que la colonne \"code_postal\" contient 35 000 valeurs "
             "distinctes ne perd pas de temps à filtrer par ville.", bold=False)

    add_para(doc, "Détection de colonnes pivot : identifier automatiquement les CODE_INSEE, "
             "SIRET, codes postaux dans chaque ressource. Un agent qui cherche des "
             "données sur Dijon reçoit directement \"filtre CODE_INSEE = 21231\".")

    add_para(doc, "Jointures cross-datasets : quand un agent consulte un dataset, suggérer "
             "les datasets qui partagent les mêmes colonnes pivot. L'open data devient "
             "un réseau, pas une collection de fichiers isolés.")

    # ── Conclusion + À propos ──

    add_heading(doc, "Conclusion", level=2)

    add_para(doc, "L'open data français est un patrimoine considérable : 75 000 jeux de données, "
             "400 000 ressources, des milliers de producteurs. Les fondations sont solides et les "
             "usages sont réels - 813 millions de téléchargements en témoignent.")

    add_para(doc, "Mais ce patrimoine est fragile. Un catalogue qui croît sans que la maintenance "
             "suive accumule de la dette technique. Des liens qui meurent, des promesses de mise à "
             "jour non tenues, des organisations qui publient puis disparaissent : autant de signaux "
             "que la quantité ne suffit pas.")

    add_para(doc, "La bonne nouvelle, c'est que les solutions sont connues et à portée de main. Un "
             "monitoring automatique de la santé des liens, un accompagnement des producteurs en "
             "difficulté, une distinction claire entre ressources publiques et internes : ces trois "
             "actions suffiraient à transformer significativement la qualité du catalogue.")

    add_para(doc, "Ce rapport est une photographie. Il appartient à chacun - plateforme, producteurs, "
             "utilisateurs - d'en faire un point de départ.")

    doc.add_page_break()

    add_heading(doc, "À propos", level=2)

    add_para(doc, "Ce rapport est le bilan du projet FlowDataGouv, un projet personnel open "
             "source mené sur quelques semaines avec des moyens limités : un serveur OVH, "
             "des agents IA, de la curiosité.")

    add_para(doc, "Les résultats présentés sont factuels et reproductibles, mais ils restent "
             "le fruit d'un audit automatisé ponctuel. Certains taux de disponibilité "
             "peuvent varier selon les conditions réseau, la charge des serveurs au "
             "moment du scan, ou des migrations en cours chez les producteurs. Ce rapport "
             "est une photographie datée du 28 mars 2026, pas un verdict définitif.")

    add_para(doc, "FlowDataGouv (demo-fli.fr) offre un accès enrichi au catalogue de "
             "data.gouv.fr : recherche sémantique, catégorisation automatique par "
             "Mistral AI, exploration tabulaire des données, et un agent conversationnel "
             "capable d'interroger les données en langage naturel.")

    add_para(doc, "L'ensemble du code source, des données d'audit et des scripts d'analyse "
             "sont librement accessibles. Toute correction, amélioration ou réutilisation "
             "est la bienvenue.")

    add_para(doc, "")
    add_para(doc, "Contact : linkedin.com/in/guillaume-clement-erp-cloud", italic=True, color=GRAY)
    add_para(doc, "Projet : demo-fli.fr", italic=True, color=GRAY)
    add_para(doc, "Code source : github.com/FLI-GCT/FlowDataGouv", italic=True, color=GRAY)

    # ── Sauvegarde ──
    output_path = OUTPUT_DIR / "rapport-datagouv-2026-publish.docx"
    doc.save(str(output_path))
    return output_path


def main():
    print("=" * 60)
    print("  03 - GÉNÉRATION DU RAPPORT WORD")
    print("=" * 60)

    output = build_report()
    print(f"\n  Rapport généré : {output}")
    print(f"  Taille : {output.stat().st_size / 1024:.0f} Ko")
    figs = list(FIGURES_DIR.glob("*.png"))
    print(f"  Figures utilisées : {len(figs)}")
    print("\n  OK - Rapport terminé")


if __name__ == "__main__":
    main()
